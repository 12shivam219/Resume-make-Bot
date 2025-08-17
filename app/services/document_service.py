from docx import Document
from docx.shared import Pt
import pythoncom
import subprocess
import logging
import os
from typing import List, Dict

logger = logging.getLogger(__name__)

class DocumentService:
    @staticmethod
    def inject_bullet_points(
        input_path: str,
        output_path: str,
        tech_stacks: Dict[str, List[str]],
        anchor_method: str = 'placeholders'
    ) -> str:
        """Main method to process resume with bullet points"""
        if input_path.lower().endswith('.doc'):
            input_path = DocumentService._convert_doc_to_docx(input_path)
        
        doc = Document(input_path)
        
        for project_num in range(1, 4):
            bullets = DocumentService._collect_bullets_for_project(tech_stacks, project_num)
            DocumentService._insert_bullets(doc, project_num, bullets, anchor_method)
        
        doc.save(output_path)
        return output_path

    @staticmethod
    def _convert_doc_to_docx(input_path: str) -> str:
        """Convert .doc to .docx using LibreOffice"""
        output_path = os.path.splitext(input_path)[0] + '.docx'
        try:
            pythoncom.CoInitialize()
            subprocess.run([
                'soffice',
                '--headless',
                '--convert-to',
                'docx',
                '--outdir',
                os.path.dirname(input_path),
                input_path
            ], check=True)
            return output_path
        except Exception as e:
            logger.error(f"DOC to DOCX conversion failed: {e}")
            raise

    @staticmethod
    def _collect_bullets_for_project(
        tech_stacks: Dict[str, List[str]],
        project_num: int
    ) -> List[str]:
        """Collect 2 bullets from each tech stack for the project"""
        bullets = []
        for stack_bullets in tech_stacks.values():
            start_idx = (project_num - 1) * 2
            bullets.extend(stack_bullets[start_idx:start_idx+2])
        return bullets

    @staticmethod
    def _insert_bullets(
        doc,
        project_num: int,
        bullets: List[str],
        anchor_method: str
    ):
        """Insert bullets using specified anchor method"""
        if anchor_method == 'placeholders':
            DocumentService._insert_via_placeholders(doc, project_num, bullets)
        else:
            DocumentService._insert_via_heuristics(doc, project_num, bullets)

    @staticmethod
    def _insert_via_placeholders(doc, project_num: int, bullets: List[str]):
        """Insert using {{PROJECTX_RESP}} placeholders"""
        placeholder = f'{{{{PROJECT{project_num}_RESP}}}}'
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, '')
                for bullet in bullets:
                    new_para = paragraph.insert_paragraph_before(bullet)
                    DocumentService._copy_paragraph_style(paragraph, new_para)

    @staticmethod
    def _insert_via_heuristics(doc, project_num: int, bullets: List[str]):
        """Insert by finding project sections"""
        project_header = f'Project {project_num}'
        for i, para in enumerate(doc.paragraphs):
            if project_header in para.text:
                for j in range(i, min(i+5, len(doc.paragraphs))):
                    if 'Responsibilities' in doc.paragraphs[j].text:
                        for bullet in reversed(bullets):
                            new_para = doc.paragraphs[j].insert_paragraph_before(bullet)
                            DocumentService._copy_paragraph_style(para, new_para)
                        return

    @staticmethod
    def _copy_paragraph_style(source_para, target_para):
        """Copy formatting from source to target paragraph"""
        target_para.style = source_para.style
        target_para.paragraph_format.left_indent = Pt(18)
        if source_para.runs and target_para.runs:
            target_para.runs[0].font.size = source_para.runs[0].font.size
            target_para.runs[0].font.name = source_para.runs[0].font.name