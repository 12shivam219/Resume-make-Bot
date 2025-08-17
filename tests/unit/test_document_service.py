import pytest
from unittest.mock import patch, MagicMock
from app.services.document_service import DocumentService
from docx import Document
import os

class TestDocumentService:
    @pytest.fixture
    def sample_docx(self, tmp_path):
        doc = Document()
        doc.add_paragraph("{{PROJECT1_RESP}}")
        doc.add_paragraph("{{PROJECT2_RESP}}")
        doc.add_paragraph("{{PROJECT3_RESP}}")
        file_path = os.path.join(tmp_path, "test.docx")
        doc.save(file_path)
        return file_path

    def test_inject_bullet_points_with_placeholders(self, sample_docx, tmp_path):
        output_path = os.path.join(tmp_path, "output.docx")
        tech_stacks = {
            "Python": [f"bullet{i}" for i in range(1, 7)],
            "Flask": [f"bullet{i}" for i in range(1, 7)],
            "AWS": [f"bullet{i}" for i in range(1, 7)]
        }
        
        result = DocumentService.inject_bullet_points(
            sample_docx,
            output_path,
            tech_stacks,
            "placeholders"
        )
        
        assert os.path.exists(result)
        doc = Document(result)
        paragraphs = [p.text for p in doc.paragraphs]
        assert "bullet1" in paragraphs
        assert "bullet2" in paragraphs

    @patch('subprocess.run')
    def test_convert_doc_to_docx(self, mock_run, tmp_path):
        mock_run.return_value = MagicMock(returncode=0)
        input_path = os.path.join(tmp_path, "test.doc")
        open(input_path, 'w').close()  # Create empty file
        
        with patch('pythoncom.CoInitialize'):
            result = DocumentService._convert_doc_to_docx(input_path)
        
        assert result.endswith(".docx")
        mock_run.assert_called_once()

    def test_collect_bullets_for_project(self):
        tech_stacks = {
            "Python": [f"bullet{i}" for i in range(1, 7)],
            "Flask": [f"bullet{i}" for i in range(1, 7)]
        }
        
        # Project 1 should get bullets 1-2 from each stack
        result = DocumentService._collect_bullets_for_project(tech_stacks, 1)
        assert result == ["bullet1", "bullet2", "bullet1", "bullet2"]
        
        # Project 2 should get bullets 3-4 from each stack
        result = DocumentService._collect_bullets_for_project(tech_stacks, 2)
        assert result == ["bullet3", "bullet4", "bullet3", "bullet4"]